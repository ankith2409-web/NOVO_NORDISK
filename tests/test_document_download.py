"""Getting the generated document out of the web interface.

Until this existed the BRD and FRD could be read on screen or regenerated from
a terminal, but the interface could not hand you the one artefact the whole
project exists to produce. That is the gap these tests cover.

Two properties matter more than the plumbing. The bytes the browser downloads
have to be the bytes ``concordance document`` writes -- a second renderer would
be a second thing to keep true, and a circulated document that differs from the
one on disk is exactly the drift this project was built to catch. And the
filename is interpolated into a quoted HTTP header, so a model name containing
a quote or a newline must not be able to close that quote or start a second
header; that is asserted directly rather than assumed from the sanitiser's
shape.
"""

from __future__ import annotations

import http.client
import threading
from http.server import ThreadingHTTPServer
from pathlib import Path

import pytest

from concordance.adapters.tmdl import TmdlAdapter
from concordance.generate import document, word
from concordance.graph.csg import SemanticGraph
from concordance.llm.fake import FakeProvider
from concordance.web import api
from concordance.web.server import _download_name, make_handler

MODEL = Path("data/models/ClinicalTrialSafety.SemanticModel")
SECOND = Path("data/models/QualityControl.SemanticModel")

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def graph() -> SemanticGraph:
    if not MODEL.exists():
        pytest.skip(f"model not present: {MODEL}")
    return SemanticGraph(TmdlAdapter().extract(str(MODEL)))


class _Server:
    def __init__(self, graph, context=None) -> None:
        handler, _ = make_handler(
            graph, FakeProvider(script=[]), context or api.ApiContext(graph=graph)
        )
        self.httpd = ThreadingHTTPServer(("127.0.0.1", 0), handler)
        self.port = self.httpd.server_port
        threading.Thread(target=self.httpd.serve_forever, daemon=True).start()

    def get(self, path: str):
        conn = http.client.HTTPConnection("127.0.0.1", self.port, timeout=10)
        conn.request("GET", path)
        return conn.getresponse()

    def close(self) -> None:
        self.httpd.shutdown()
        self.httpd.server_close()


@pytest.fixture
def server(graph):
    running = _Server(graph)
    yield running
    running.close()


# -- the download itself ------------------------------------------------------

def test_the_business_document_downloads_as_markdown(server) -> None:
    response = server.get("/api/document?kind=business")
    body = response.read().decode("utf-8")

    assert response.status == 200
    assert "text/markdown" in response.getheader("Content-Type")
    assert body.startswith("# Business Requirements Document")


def test_it_is_sent_as_an_attachment_with_a_descriptive_name(server) -> None:
    """`attachment` is what makes a browser save rather than render it, and the
    name is what stops three models' documents becoming `document (2).md`."""
    response = server.get("/api/document?kind=business")
    response.read()

    disposition = response.getheader("Content-Disposition")
    assert disposition == 'attachment; filename="ClinicalTrialSafety-BRD.md"'


def test_the_functional_document_is_a_different_document(server) -> None:
    brd = server.get("/api/document?kind=business").read().decode()
    frd = server.get("/api/document?kind=functional").read().decode()

    assert frd.startswith("# Functional Requirements Document")
    assert frd != brd


def test_word_format_returns_a_real_docx(server) -> None:
    response = server.get("/api/document?kind=business&format=docx")
    body = response.read()

    assert response.status == 200
    assert response.getheader("Content-Type") == DOCX_TYPE
    assert 'filename="ClinicalTrialSafety-BRD.docx"' in response.getheader(
        "Content-Disposition"
    )
    # A .docx is a zip; the magic number is the cheapest proof this is one
    # rather than an error page with an optimistic content type.
    assert body[:2] == b"PK"


def test_the_downloaded_bytes_are_what_the_cli_would_write(server, graph) -> None:
    """The property that matters most.

    The interface and the command line must not be able to disagree about what
    the document says, so this asserts they come from one renderer rather than
    trusting that they do.
    """
    downloaded = server.get("/api/document?kind=business").read().decode("utf-8")
    rendered = document.to_markdown(document.build(graph, document.Kind.BUSINESS))

    assert downloaded == rendered


def test_the_word_bytes_match_the_same_document(graph) -> None:
    """`to_bytes` and `write` must render identically -- they are the same call.

    Compared through python-docx's own reading rather than byte-for-byte: a
    .docx is a zip, and two zips of identical content differ in their embedded
    timestamps.
    """
    import io

    import docx

    built = document.build(graph, document.Kind.BUSINESS)
    from_memory = docx.Document(io.BytesIO(word.to_bytes(built)))
    reference = word.render(built)

    assert [p.text for p in from_memory.paragraphs] == [
        p.text for p in reference.paragraphs
    ]


# -- selecting what to download ----------------------------------------------

def test_each_model_downloads_its_own_document(graph) -> None:
    if not SECOND.exists():
        pytest.skip(f"model not present: {SECOND}")
    other = SemanticGraph(TmdlAdapter().extract(str(SECOND)))
    registry = api.ModelRegistry(
        contexts={
            graph.model.name: api.ApiContext(graph=graph),
            other.model.name: api.ApiContext(graph=other),
        },
        default=graph.model.name,
    )
    running = _Server(graph, context=registry)
    try:
        response = running.get(f"/api/document?kind=business&model={other.model.name}")
        body = response.read().decode()
        assert f'filename="{other.model.name}-BRD.md"' in response.getheader(
            "Content-Disposition"
        )
        assert other.model.name in body
    finally:
        running.close()


def test_an_unknown_model_is_refused_without_echoing_it_back(server) -> None:
    response = server.get("/api/document?kind=business&model=not-loaded-here")
    body = response.read().decode()

    assert response.status == 404
    assert "not-loaded-here" not in body


def test_an_unknown_kind_is_a_readable_refusal(server) -> None:
    response = server.get("/api/document?kind=marketing")
    body = response.read().decode()

    assert response.status == 400
    assert "business" in body and "functional" in body


def test_an_unknown_format_is_a_readable_refusal(server) -> None:
    response = server.get("/api/document?kind=business&format=pdf")
    body = response.read().decode()

    assert response.status == 400
    assert "md" in body and "docx" in body


def test_the_default_is_the_business_document(server) -> None:
    """A bare /api/document must not 500 or guess unhelpfully."""
    response = server.get("/api/document")
    body = response.read().decode()

    assert response.status == 200
    assert body.startswith("# Business Requirements Document")


# -- the filename cannot forge a header ---------------------------------------

def test_a_quote_in_a_model_name_cannot_close_the_header_value() -> None:
    """The filename lands inside `filename="..."`. A model name that could
    close that quote would let the rest of it be read as header directives."""
    name = _download_name('evil"; download-everything; x="', "business")

    assert '"' not in name
    assert ";" not in name


def test_a_newline_in_a_model_name_cannot_start_a_second_header() -> None:
    """CRLF injection: the classic way a header value becomes two headers."""
    name = _download_name("model\r\nX-Injected: yes", "business")

    assert "\r" not in name
    assert "\n" not in name
    assert ":" not in name


def test_path_separators_never_survive_into_a_filename() -> None:
    """A browser is told what to call the file; `../` in that name is a request
    to write outside the downloads folder."""
    name = _download_name("../../etc/passwd", "business")

    assert "/" not in name
    assert name.startswith("..-..-etc-passwd")


def test_a_name_with_nothing_usable_still_yields_a_filename() -> None:
    """Not a crash, and not an empty `filename=""` the browser has to invent
    a name for."""
    assert _download_name("///", "business") == "model-BRD.md"
    assert _download_name("", "functional") == "model-FRD.md"


def test_the_extension_follows_the_requested_format() -> None:
    assert _download_name("Sales", "business", "md").endswith("-BRD.md")
    assert _download_name("Sales", "functional", "docx").endswith("-FRD.docx")


def test_the_download_route_is_listed_where_routes_are_listed(server) -> None:
    """A route the 404 body omits is a route nobody finds.

    It cannot live in `api.ROUTES` -- everything there answers with JSON and
    this answers with a document -- so the risk is exactly that it gets counted
    nowhere and quietly becomes undiscoverable.
    """
    response = server.get("/api/nonsense")
    body = response.read().decode()

    assert response.status == 404
    assert "/api/document" in body
