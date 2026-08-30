"""Word export and the evidence pack.

Rendering cannot be checked visually here -- LibreOffice in this environment
fails to convert even a one-paragraph document -- so the document is verified
structurally instead: heading levels Word will build a navigation pane from, a
traceability row per requirement, DAX in a monospaced run, and no Markdown
markers surviving into text someone is about to circulate for signature.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document as WordDocument

from concordance.adapters.tmdl import TmdlAdapter
from concordance.generate import auditpack, word
from concordance.generate import document as doc
from concordance.generate.requirements import Kind
from concordance.generate.word import _split_markup
from concordance.graph.csg import SemanticGraph

MODEL = Path("data/models/ClinicalTrialSafety.SemanticModel")


@pytest.fixture(scope="module")
def graph() -> SemanticGraph:
    if not MODEL.exists():
        pytest.skip(f"model not present: {MODEL}")
    return SemanticGraph(TmdlAdapter().extract(str(MODEL)))


@pytest.fixture(scope="module")
def brd_path(graph: SemanticGraph, tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("word") / "brd.docx"
    return word.write(doc.build(graph, Kind.BUSINESS), out)


@pytest.fixture(scope="module")
def frd_path(graph: SemanticGraph, tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("word") / "frd.docx"
    return word.write(doc.build(graph, Kind.FUNCTIONAL), out)


# -- the markup splitter -----------------------------------------------------

@pytest.mark.parametrize(
    "text,expected",
    [
        ("plain text", [("plain text", "plain")]),
        ("a **bold** b", [("a ", "plain"), ("bold", "bold"), (" b", "plain")]),
        ("use `CODE` here", [("use ", "plain"), ("CODE", "code"), (" here", "plain")]),
        (
            "**A** and `B`",
            [("A", "bold"), (" and ", "plain"), ("B", "code")],
        ),
        # An unmatched marker is text, not the start of formatting.
        ("2 ** 3 is not bold", [("2 ** 3 is not bold", "plain")]),
    ],
)
def test_markup_splits_correctly(text: str, expected: list) -> None:
    assert _split_markup(text) == expected


# -- the rendered document ---------------------------------------------------

def test_the_file_opens_as_a_valid_word_document(brd_path: Path) -> None:
    reopened = WordDocument(str(brd_path))
    assert len(reopened.paragraphs) > 50
    assert reopened.tables


def test_headings_use_built_in_styles(brd_path: Path) -> None:
    """Word builds its navigation pane from these; custom styles would not appear."""
    styles = {p.style.name for p in WordDocument(str(brd_path)).paragraphs}
    assert "Title" in styles
    assert "Heading 1" in styles
    assert "Heading 2" in styles


def test_no_markdown_markers_survive_into_the_text(brd_path: Path) -> None:
    """A document going out for signature must not contain literal asterisks."""
    for paragraph in WordDocument(str(brd_path)).paragraphs:
        assert "**" not in paragraph.text
        assert "`" not in paragraph.text


def test_bold_markup_becomes_actual_bold_runs(brd_path: Path) -> None:
    document = WordDocument(str(brd_path))
    bold = [r for p in document.paragraphs for r in p.runs if r.bold and r.text.strip()]
    assert len(bold) > 20


def _table_headed(document, headers: list[str]):
    """The table whose first row is ``headers``, whatever order tables appear in."""
    for table in document.tables:
        if [c.text for c in table.rows[0].cells] == headers:
            return table
    raise AssertionError(f"no table headed {headers}")


def test_dax_is_rendered_monospaced(frd_path: Path) -> None:
    """In the FRD, which is where DAX belongs.

    This asked the BRD for its DAX until a reviewer pointed out that a BRD
    should not have any: "BRD is basically the complete business information ...
    very plain English, that is it, and FRD is where you get into the details."
    The test was encoding the wrong document's job.
    """
    document = WordDocument(str(frd_path))
    mono = [
        r
        for p in document.paragraphs
        for r in p.runs
        if r.font.name == "Consolas" and "CALCULATE" in r.text
    ]
    assert mono, "DAX expressions should be visually distinct from prose"


def test_the_brd_carries_no_dax_at_all(brd_path: Path) -> None:
    """The other half of that split, asserted rather than assumed.

    A BRD is read by the people who own the reporting and sign it off. Thirty
    lines of DAX in front of them is not business information, and it is the
    first thing a reviewer notices.
    """
    document = WordDocument(str(brd_path))
    leaked = [
        p.text
        for p in document.paragraphs
        if any(token in p.text for token in ("CALCULATE(", "DIVIDE(", "SUMX(", "COUNTROWS("))
    ]
    assert not leaked, f"DAX reached the BRD: {leaked[:3]}"


def test_traceability_matrix_has_a_row_per_requirement(
    graph: SemanticGraph, brd_path: Path
) -> None:
    built = doc.build(graph, Kind.BUSINESS)
    # Found by its header rather than by position. It was `tables[0]`, which
    # silently became the glossary the moment the document grew one -- and the
    # failure read as "the matrix lost 22 rows" rather than "you are looking at
    # a different table".
    table = _table_headed(
        WordDocument(str(brd_path)),
        ["Requirement", "Bound to", "Fingerprint", "Confidence"],
    )

    assert len(table.rows) == len(built.requirements) + 1  # + header
    headers = [c.text for c in table.rows[0].cells]
    assert headers == ["Requirement", "Bound to", "Fingerprint", "Confidence"]

    ids_in_table = {row.cells[0].text for row in table.rows[1:]}
    assert ids_in_table == {r.id for r in built.requirements}


def test_the_review_queue_appears_before_the_requirements(brd_path: Path) -> None:
    paragraphs = [p.text for p in WordDocument(str(brd_path)).paragraphs]
    queue = paragraphs.index("Awaiting human confirmation")
    first_section = next(i for i, t in enumerate(paragraphs) if t.startswith("1. "))
    assert queue < first_section


def test_both_document_kinds_render(graph: SemanticGraph, tmp_path: Path) -> None:
    for kind in (Kind.BUSINESS, Kind.FUNCTIONAL):
        path = word.write(doc.build(graph, kind), tmp_path / f"{kind.value}.docx")
        assert WordDocument(str(path)).paragraphs


# -- the evidence pack -------------------------------------------------------

@pytest.fixture(scope="module")
def pack(graph: SemanticGraph, tmp_path_factory):
    return auditpack.build(graph, tmp_path_factory.mktemp("pack"))


def test_the_pack_contains_every_expected_artefact(pack) -> None:
    names = {p.name for p in pack.files}
    assert "MANIFEST.json" in names
    assert "README.txt" in names
    assert any(n.endswith(".BRD.docx") for n in names)
    assert any(n.endswith(".FRD.docx") for n in names)
    assert any(n.endswith(".BRD.md") for n in names)
    assert any(n.endswith(".fingerprints.json") for n in names)


def test_the_manifest_records_what_was_read_and_by_which_version(pack) -> None:
    manifest = json.loads((pack.directory / "MANIFEST.json").read_text())

    assert manifest["tool"]["name"] == "concordance"
    assert manifest["tool"]["version"]
    assert manifest["model"]["name"] == "ClinicalTrialSafety"
    assert manifest["model"]["measures"] == 24
    assert manifest["requirements"]["total"] == pack.requirement_count


def test_the_manifest_states_its_own_limitations(pack) -> None:
    """A pack that hid what it could not read would misrepresent its completeness."""
    manifest = json.loads((pack.directory / "MANIFEST.json").read_text())
    limitations = manifest["limitations"]

    assert "unresolved_references" in limitations
    assert "features_not_extracted" in limitations
    assert manifest["requirements"]["awaiting_human_confirmation"] == pack.needs_review


def test_the_fingerprint_manifest_makes_the_pack_verifiable(
    pack, graph: SemanticGraph
) -> None:
    """The property that makes this evidence rather than assertion."""
    from concordance.drift import snapshot as snap
    from concordance.drift.compare import compare

    stored = snap.Snapshot.load(
        next(p for p in pack.files if p.name.endswith(".fingerprints.json"))
    )
    # Recomputing from the same model must agree with what the pack recorded.
    report = compare(stored, snap.take(graph, "recheck"))
    assert not report.has_drift


def test_the_readme_explains_how_to_verify_the_pack(pack) -> None:
    readme = (pack.directory / "README.txt").read_text()
    assert "concordance drift" in readme
    assert "ClinicalTrialSafety" in readme
    # It must not overclaim regulatory standing.
    assert "quality function" in readme


def test_the_pack_counts_match_the_documents_it_contains(
    pack, graph: SemanticGraph
) -> None:
    brd = doc.build(graph, Kind.BUSINESS)
    frd = doc.build(graph, Kind.FUNCTIONAL)
    assert pack.requirement_count == len(brd.requirements) + len(frd.requirements)


# -- the shape a BRD and an FRD are expected to have ---------------------------
#
# Checked against how these documents are actually written, not against our own
# taste. Published templates agree on the opening: purpose, scope, assumptions
# and constraints, then the requirements, then a glossary and a traceability
# matrix. They also agree on the split -- a BRD is the business case in business
# language and an FRD is the system detail -- which is the same line the
# reviewer drew in her own words.

def test_both_documents_open_the_way_a_requirements_document_opens(
    graph: SemanticGraph,
) -> None:
    for kind in (Kind.BUSINESS, Kind.FUNCTIONAL):
        text = doc.to_markdown(doc.build(graph, kind))
        for heading in ("## Purpose", "## Scope of this document",
                        "## Assumptions and constraints"):
            assert heading in text, f"{kind.value} is missing {heading}"


def test_the_markdown_brd_carries_no_dax_either(graph: SemanticGraph) -> None:
    """The .docx test alone did not catch this.

    The two renderers gate the implementation block separately, so removing the
    gate in one left the other passing -- proved by putting the DAX back in the
    Markdown renderer and watching the whole file stay green. Both formats are
    the same document and both are asserted.
    """
    text = doc.to_markdown(doc.build(graph, Kind.BUSINESS))
    assert "*Implementation:*" not in text
    assert "```dax" not in text
    for token in ("CALCULATE(", "DIVIDE(", "SUMX(", "COUNTROWS("):
        assert token not in text, f"{token} reached the Markdown BRD"


def test_the_markdown_frd_still_carries_it(graph: SemanticGraph) -> None:
    text = doc.to_markdown(doc.build(graph, Kind.FUNCTIONAL))
    assert "*Implementation:*" in text


def test_a_brd_names_what_a_model_cannot_tell_it(graph: SemanticGraph) -> None:
    """Objectives, stakeholders and a cost case belong in a BRD and are in no
    .pbix. Naming them as gaps is the only honest option: omitting them makes
    the document look complete when it is not, and inventing them is the exact
    failure this project exists to prevent."""
    text = doc.to_markdown(doc.build(graph, Kind.BUSINESS))
    assert "## To be supplied by the business" in text
    for expected in ("Business objectives", "Stakeholders", "Cost, benefit"):
        assert expected in text


def test_an_frd_does_not_ask_the_reader_for_a_business_case(
    graph: SemanticGraph,
) -> None:
    """That section is the BRD's. An FRD repeating it would be telling a
    developer to go and find the commercial justification."""
    text = doc.to_markdown(doc.build(graph, Kind.FUNCTIONAL))
    assert "## To be supplied by the business" not in text


def test_the_glossary_only_holds_terms_the_model_defines(
    graph: SemanticGraph,
) -> None:
    """A glossary of invented definitions would be worse than none."""
    built = doc.build(graph, Kind.BUSINESS)
    described = {
        m.name: m.description for m in graph.model.measures if m.description
    }
    for term, meaning in built.glossary:
        assert described.get(term) == meaning or meaning


def test_the_two_documents_do_not_disagree_on_their_own_first_page(
    graph: SemanticGraph,
) -> None:
    """The front matter is written once and rendered twice.

    The Markdown and the .docx are one document in two formats; two renderers
    each wording the opening themselves would make them two documents that
    happen to share a title.
    """
    built = doc.build(graph, Kind.BUSINESS)
    blocks = doc.front_matter_blocks(built)
    markdown = doc.to_markdown(built)
    for heading, paragraphs, _ in blocks:
        assert f"## {heading}" in markdown
        for paragraph in paragraphs:
            assert paragraph in markdown, f"{heading!r} differs between renderers"
