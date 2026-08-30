"""Word export for generated requirements documents.

A BRD in Markdown is a developer artefact. The people who sign one off open it
in Word, comment in the margin, and route it for approval -- and the hackathon's
own rules require Office 365 "for the exchange of material", which is a fair
signal of how these documents actually travel.

``python-docx`` rather than the Node ``docx`` library, deliberately: the project
installs with pip and must keep doing so. Requiring Node alongside it to render
one output format would be a poor trade for a team cloning this repo.

Nothing here decides content. The document is rendered from the same
``Document`` object the Markdown writer uses, so the two formats cannot drift
apart in what they claim.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from docx.shared import Pt, RGBColor

from concordance.generate.document import Document
from concordance.generate.requirements import Confidence

#: Confidence shown in the colour a reader already associates with it, so a
#: page of requirements can be skimmed for what still needs attention.
_CONFIDENCE_COLOUR = {
    Confidence.HIGH: RGBColor(0x2E, 0x7D, 0x5B),
    Confidence.MEDIUM: RGBColor(0x9A, 0x6B, 0x1E),
    Confidence.LOW: RGBColor(0xA3, 0x3B, 0x3B),
}

_MONO = "Consolas"


def render(document: Document) -> WordDocument:
    """Build the Word document in memory, without deciding where it goes.

    Split out from ``write`` so the web server can serve the same bytes over a
    socket that the CLI writes to a path. Rendering twice -- once for each
    destination -- would be two things to keep true, and a circulated document
    that differs from the one on disk is precisely the drift this project
    exists to catch.
    """
    word = WordDocument()
    _configure_styles(word)

    _title_block(word, document)
    _front_matter(word, document)
    _review_queue(word, document)
    _requirements(word, document)
    _glossary(word, document)
    _traceability_matrix(word, document)
    return word


def to_bytes(document: Document) -> bytes:
    """The same .docx, as bytes, for callers with no filesystem to write to."""
    buffer = BytesIO()
    render(document).save(buffer)
    return buffer.getvalue()


def write(document: Document, path: str | Path) -> Path:
    """Render a BRD or FRD to a .docx file."""
    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    render(document).save(str(destination))
    return destination


# -- sections ----------------------------------------------------------------

def _configure_styles(word: WordDocument) -> None:
    normal = word.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def _title_block(word: WordDocument, document: Document) -> None:
    word.add_heading(document.title, level=0)

    counts = document.counts()
    meta = word.add_paragraph()
    meta.add_run("Source model: ").bold = True
    meta.add_run(document.source).font.name = _MONO
    meta.add_run("\nGenerated: ").bold = True
    meta.add_run(document.generated_on)
    meta.add_run("\nRequirements: ").bold = True
    meta.add_run(
        f"{counts['requirements']} "
        f"({counts['high']} stated, {counts['medium']} inferred, "
        f"{counts['low']} need confirmation)"
    )

    note = word.add_paragraph()
    note.paragraph_format.left_indent = Pt(18)
    run = note.add_run(
        "Every requirement below was derived from the implemented semantic model "
        "and is bound to the fingerprint of the object that satisfies it. "
        "Confidence reflects whether the model states the requirement or whether "
        "it was inferred from structure. Inferred items are marked and must be "
        "confirmed by a human before this document is treated as approved."
    )
    run.italic = True
    run.font.size = Pt(9.5)


def _review_queue(word: WordDocument, document: Document) -> None:
    if not document.review_queue:
        return

    word.add_heading("Awaiting human confirmation", level=1)
    word.add_paragraph(
        f"{len(document.review_queue)} requirement(s) could not be established "
        f"from the model alone and need a decision before sign-off:"
    )
    for requirement in document.review_queue:
        item = word.add_paragraph(style="List Bullet")
        item.add_run(f"{requirement.id} — ").bold = True
        item.add_run(_plain(requirement.statement))


def _requirements(word: WordDocument, document: Document) -> None:
    for index, section in enumerate(document.sections, start=1):
        word.add_heading(f"{index}. {section.title}", level=1)

        for requirement in section.requirements:
            heading = word.add_heading(requirement.id, level=2)
            if requirement.needs_review:
                flag = heading.add_run("  — needs confirmation")
                flag.italic = True
                flag.font.color.rgb = _CONFIDENCE_COLOUR[Confidence.LOW]

            _rich_paragraph(word, requirement.statement)

            rationale = word.add_paragraph()
            label = rationale.add_run("Rationale: ")
            label.bold = True
            label.font.size = Pt(9.5)
            detail = rationale.add_run(requirement.rationale)
            detail.font.size = Pt(9.5)

            confidence = word.add_paragraph()
            confidence_label = confidence.add_run("Confidence: ")
            confidence_label.bold = True
            confidence_label.font.size = Pt(9.5)
            value = confidence.add_run(requirement.confidence.value)
            value.font.size = Pt(9.5)
            value.font.color.rgb = _CONFIDENCE_COLOUR[requirement.confidence]

            _implementation(word, requirement, document)
            _sql(word, document, requirement)


def _sql(word: WordDocument, document: Document, requirement) -> None:
    """The same measure as SQL, directly beneath its DAX.

    Kept beside the requirement rather than gathered into an appendix, for the
    same reason as in the Markdown: a reader who has found the requirement has
    found the query, without a second lookup.
    """
    from concordance.generate.document import measure_of

    translation = document.sql.get(measure_of(requirement))
    if translation is None:
        return

    query = getattr(translation, "sql", "")
    if query:
        caption = word.add_paragraph()
        run = caption.add_run(f"Equivalent SQL ({document.sql_dialect})")
        run.bold = True
        run.font.size = Pt(9.5)

        code = word.add_paragraph()
        code.paragraph_format.left_indent = Pt(18)
        # One run per line: Word does not break a run on a newline, so a single
        # run would render the whole query as one unwrapped line.
        for number, line in enumerate(query.splitlines()):
            text = code.add_run(("\n" if number else "") + line)
            text.font.name = _MONO
            text.font.size = Pt(8.5)
        return

    note = word.add_paragraph()
    reason = (getattr(translation, "reason", "") or "").rstrip(".")
    run = note.add_run(
        f"No SQL equivalent: {reason}. This is a property of the expression "
        "rather than a gap in the translation."
    )
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = _CONFIDENCE_COLOUR[Confidence.LOW]


def _front_matter(word: WordDocument, document: Document) -> None:
    """Purpose, scope, constraints -- the sections a real BRD or FRD opens with.

    Rendered from the same text the Markdown renderer uses, because two
    renderers writing the same sections in their own words is two documents,
    and the whole premise here is that the .docx and the .md are one document
    in two formats.
    """
    from concordance.generate.document import front_matter_blocks

    for heading, paragraphs, bullets in front_matter_blocks(document):
        word.add_heading(heading, level=1)
        for text in paragraphs:
            word.add_paragraph(text)
        for text in bullets:
            word.add_paragraph(text, style="List Bullet")


def _glossary(word: WordDocument, document: Document) -> None:
    if not document.glossary:
        return
    word.add_heading("Glossary", level=1)
    word.add_paragraph(
        "Taken from the descriptions recorded in the model. Terms with no "
        "description there do not appear here rather than being given one."
    )
    table = word.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Term"
    header[1].text = "Meaning"
    for term, meaning in document.glossary:
        row = table.add_row().cells
        row[0].text = term
        row[1].text = meaning


def _implementation(word: WordDocument, requirement, document: Document) -> None:
    """Show the DAX only where a single piece of evidence is the implementation.

    And only in the FRD: the reviewer's own split is that a BRD is business
    information in plain English and the FRD is where the detail lives.
    """
    from concordance.generate.requirements import Kind

    if document.kind is not Kind.FUNCTIONAL:
        if len(requirement.evidence) > 1:
            note = word.add_paragraph()
            run = note.add_run(
                f"Bound to {len(requirement.evidence)} objects — see the "
                f"traceability matrix."
            )
            run.italic = True
            run.font.size = Pt(9.5)
        return
    if len(requirement.evidence) == 1 and requirement.evidence[0].detail:
        caption = word.add_paragraph()
        run = caption.add_run("Implementation")
        run.bold = True
        run.font.size = Pt(9.5)

        code = word.add_paragraph()
        code.paragraph_format.left_indent = Pt(18)
        expression = code.add_run(requirement.evidence[0].detail)
        expression.font.name = _MONO
        expression.font.size = Pt(9)
    elif len(requirement.evidence) > 1:
        note = word.add_paragraph()
        run = note.add_run(
            f"Bound to {len(requirement.evidence)} objects — see the traceability matrix."
        )
        run.italic = True
        run.font.size = Pt(9.5)


def _traceability_matrix(word: WordDocument, document: Document) -> None:
    word.add_page_break()
    word.add_heading("Traceability matrix", level=1)
    word.add_paragraph(
        "Each requirement is bound to the fingerprint of the object that "
        "satisfies it. A change to that object moves its fingerprint, which is "
        "what makes drift against this document detectable rather than a matter "
        "of opinion."
    )

    table = word.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    for cell, heading in zip(
        table.rows[0].cells, ("Requirement", "Bound to", "Fingerprint", "Confidence")
    ):
        cell.text = ""
        run = cell.paragraphs[0].add_run(heading)
        run.bold = True
        run.font.size = Pt(9)

    for requirement in document.requirements:
        if requirement.evidence:
            first = requirement.evidence[0]
            target = first.node_id
            if len(requirement.evidence) > 1:
                target += f" (+{len(requirement.evidence) - 1} more)"
            fingerprint = first.fingerprint[:12]
        else:
            target = "—"
            fingerprint = "—"

        cells = table.add_row().cells
        for cell, text, mono in (
            (cells[0], requirement.id, True),
            (cells[1], target, True),
            (cells[2], fingerprint, True),
            (cells[3], requirement.confidence.value, False),
        ):
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(8)
            if mono:
                run.font.name = _MONO
            if cell is cells[3]:
                run.font.color.rgb = _CONFIDENCE_COLOUR[requirement.confidence]


# -- helpers -----------------------------------------------------------------

def _rich_paragraph(word: WordDocument, text: str) -> None:
    """Render the light Markdown used in statements: **bold** and `code`.

    Written as a small tokeniser rather than a regex replace because the two
    markers interleave, and a naive pass would emit literal asterisks into a
    document someone is about to circulate for signature.
    """
    paragraph = word.add_paragraph()
    for chunk, style in _split_markup(text):
        run = paragraph.add_run(chunk)
        if style == "bold":
            run.bold = True
        elif style == "code":
            run.font.name = _MONO
            run.font.size = Pt(9.5)


def _split_markup(text: str) -> list[tuple[str, str]]:
    """Split text into (chunk, style) where style is plain, bold or code."""
    pieces: list[tuple[str, str]] = []
    buffer: list[str] = []
    index = 0

    def flush() -> None:
        if buffer:
            pieces.append(("".join(buffer), "plain"))
            buffer.clear()

    while index < len(text):
        if text.startswith("**", index):
            closing = text.find("**", index + 2)
            if closing != -1:
                flush()
                pieces.append((text[index + 2 : closing], "bold"))
                index = closing + 2
                continue
        if text[index] == "`":
            closing = text.find("`", index + 1)
            if closing != -1:
                flush()
                pieces.append((text[index + 1 : closing], "code"))
                index = closing + 1
                continue
        buffer.append(text[index])
        index += 1

    flush()
    return pieces


def _plain(text: str) -> str:
    return text.replace("**", "").replace("`", "")
